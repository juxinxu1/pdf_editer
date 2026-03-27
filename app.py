"""
PDF Studio — Editor estructurado con control de fuente
Edita el PDF bloque a bloque, preservando posiciones y páginas.
"""

import streamlit as st
import io
import json
import zipfile
import re
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Optional

# ─── PAGE CONFIG ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PDF Studio",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─── CSS ──────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Mono:wght@300;400;500&family=Fraunces:ital,opsz,wght@0,9..144,300;0,9..144,600;1,9..144,300&display=swap');

html, body, [class*="css"] { font-family: 'DM Mono', monospace !important; }
.stApp { background-color: #f4efe6; }

.hdr { background:#111; color:#f4efe6; padding:18px 28px; border-radius:6px; margin-bottom:18px; }
.hdr h1 { font-family:'Fraunces',serif; font-size:1.75rem; font-weight:300; font-style:italic; margin:0; color:#f4efe6; }
.hdr p  { font-size:0.65rem; letter-spacing:0.14em; text-transform:uppercase; color:#aaa; margin:3px 0 0 0; }

.ok-box  { background:#e8f5e9; border-left:3px solid #388e3c; padding:9px 14px; border-radius:0 3px 3px 0; font-size:0.73rem; color:#1b5e20; margin:8px 0; }
.err-box { background:#ffebee; border-left:3px solid #c62828; padding:9px 14px; border-radius:0 3px 3px 0; font-size:0.73rem; color:#b71c1c; margin:8px 0; }
.warn-box{ background:#fff8e1; border-left:3px solid #f9a825; padding:9px 14px; border-radius:0 3px 3px 0; font-size:0.73rem; color:#6d4c00; margin:8px 0; }
.info-bar{ background:#ede7d9; border:1px solid rgba(0,0,0,0.1); border-radius:3px; padding:9px 14px; font-size:0.71rem; color:#3d4a5c; margin-bottom:10px; }

.page-divider {
  text-align:center; font-size:0.6rem; letter-spacing:0.2em; text-transform:uppercase;
  color:#aaa; border-top:1px dashed #ccc; padding-top:10px; margin:18px 0 10px 0;
}

.stButton>button {
  font-family:'DM Mono',monospace !important; font-size:0.68rem !important;
  letter-spacing:0.1em !important; text-transform:uppercase !important;
  border-radius:2px !important; border:1.5px solid #111 !important;
  background:#111 !important; color:#f4efe6 !important;
}
.stButton>button:hover { background:#333 !important; }
.stDownloadButton>button {
  font-family:'DM Mono',monospace !important; font-size:0.68rem !important;
  letter-spacing:0.1em !important; text-transform:uppercase !important;
  border-radius:2px !important; background:#c47b28 !important;
  color:white !important; border:none !important;
}
.stTextArea textarea {
  font-family:'DM Mono',monospace !important; font-size:0.8rem !important;
  background:#fdfaf5 !important; border:1.5px solid rgba(0,0,0,0.15) !important;
  border-radius:2px !important; line-height:1.75 !important;
}
.stTabs [data-baseweb="tab-list"] { background:#e8e2d9; border-bottom:1.5px solid #111; gap:0; }
.stTabs [data-baseweb="tab"] {
  font-family:'DM Mono',monospace !important; font-size:0.68rem !important;
  letter-spacing:0.1em !important; text-transform:uppercase !important;
  color:#555 !important; border-right:1px solid rgba(0,0,0,0.1) !important;
  padding:12px 20px !important;
}
.stTabs [aria-selected="true"] { background:#fff !important; color:#111 !important; border-bottom:2.5px solid #c47b28 !important; }
div[data-testid="stFileUploader"] { border:2px dashed #888; border-radius:4px; background:#fff; padding:8px; }
div[data-testid="stFileUploader"]:hover { border-color:#c47b28; }
.sec-title { font-family:'Fraunces',serif !important; font-size:1.35rem; font-weight:300; font-style:italic; color:#111; }
.sec-sub   { font-size:0.64rem; letter-spacing:0.12em; text-transform:uppercase; color:#666; margin-bottom:16px; }
</style>
""", unsafe_allow_html=True)

# ─── HEADER ───────────────────────────────────────────────────────────────────
st.markdown("""
<div class="hdr">
  <h1>● PDF Studio</h1>
  <p>Editor estructurado · Convertidor · Asistente IA</p>
</div>
""", unsafe_allow_html=True)

# ─── SESSION STATE ────────────────────────────────────────────────────────────
DEFAULTS = {
    "blocks": [],
    "pdf_bytes": None,
    "pdf_filename": "",
    "n_pages": 0,
    "ai_result": "",
    "api_key": "",
    "conv_results": [],
    "ai_prompt_v": "",
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ─── HELPER FUNCTIONS ─────────────────────────────────────────────────────────

def extract_blocks(pdf_bytes: bytes):
    """Extract text blocks preserving positions. Returns (blocks, n_pages, error)."""
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        out = []
        for pn in range(len(doc)):
            page = doc[pn]
            for b in page.get_text("blocks"):
                if b[6] != 0:
                    continue
                text = b[4].strip()
                if not text:
                    continue
                out.append({
                    "page": pn, "block_idx": b[5],
                    "original_text": text, "edited_text": text,
                    "x0": round(b[0],2), "y0": round(b[1],2),
                    "x1": round(b[2],2), "y1": round(b[3],2),
                    "font_name": "Helvetica", "font_size": 11.0,
                    "font_color": "#000000", "bold": False, "italic": False,
                })
        return out, len(doc), ""
    except ImportError:
        pass
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        out = []
        for pn, page in enumerate(reader.pages):
            raw = (page.extract_text() or "").strip()
            for bi, para in enumerate(p for p in raw.split("\n\n") if p.strip()):
                out.append({
                    "page": pn, "block_idx": bi,
                    "original_text": para, "edited_text": para,
                    "x0": 50.0, "y0": float(bi * 60),
                    "x1": 550.0, "y1": float(bi * 60 + 50),
                    "font_name": "Helvetica", "font_size": 11.0,
                    "font_color": "#000000", "bold": False, "italic": False,
                })
        return out, len(reader.pages), ""
    except Exception as e:
        return [], 0, str(e)


def hex_to_rgb(h: str):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) / 255.0 for i in (0, 2, 4))


def blocks_to_pdf(blocks: list, original_pdf_bytes: bytes) -> bytes:
    """Overlay edited text on original PDF using PyMuPDF."""
    try:
        import fitz
        doc = fitz.open(stream=original_pdf_bytes, filetype="pdf")
        by_page = {}
        for b in blocks:
            by_page.setdefault(b["page"], []).append(b)

        for pn, pblocks in by_page.items():
            if pn >= len(doc):
                continue
            page = doc[pn]
            for blk in pblocks:
                changed = (blk["edited_text"] != blk["original_text"]
                           or blk.get("bold") or blk.get("italic")
                           or blk.get("font_color", "#000000") != "#000000"
                           or blk.get("font_name", "Helvetica") != "Helvetica"
                           or blk.get("font_size", 11.0) != 11.0)
                if not changed:
                    continue

                rect = fitz.Rect(blk["x0"], blk["y0"], blk["x1"], blk["y1"])
                page.draw_rect(rect, color=(1,1,1), fill=(1,1,1))

                bold   = blk.get("bold", False)
                italic = blk.get("italic", False)
                fname  = blk.get("font_name", "Helvetica")

                fitz_fonts = {
                    "Helvetica":       {(0,0):"helv",(1,0):"hebo",(0,1):"heit",(1,1):"hebi"},
                    "Times New Roman": {(0,0):"tiro",(1,0):"tibo",(0,1):"tiit",(1,1):"tibi"},
                    "Courier":         {(0,0):"cour",(1,0):"cobo",(0,1):"coit",(1,1):"cobi"},
                }
                fitz_font = fitz_fonts.get(fname, fitz_fonts["Helvetica"]).get(
                    (int(bold), int(italic)), "helv")

                fsize  = float(blk.get("font_size", 11.0))
                r, g, b_c = hex_to_rgb(blk.get("font_color", "#000000"))

                y_cur = blk["y0"] + fsize
                for line in blk["edited_text"].split("\n"):
                    if y_cur > blk["y1"] + fsize * 5:
                        break
                    page.insert_text(
                        fitz.Point(blk["x0"], y_cur), line,
                        fontname=fitz_font, fontsize=fsize, color=(r, g, b_c),
                    )
                    y_cur += fsize * 1.35

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return _blocks_to_pdf_rl(blocks)


def _blocks_to_pdf_rl(blocks: list) -> bytes:
    """Reportlab fallback PDF builder."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    story = []
    cur_page = -1

    for blk in sorted(blocks, key=lambda b: (b["page"], b.get("y0", 0))):
        if blk["page"] != cur_page:
            if cur_page >= 0:
                story.append(PageBreak())
            cur_page = blk["page"]

        fname  = blk.get("font_name", "Helvetica")
        bold   = blk.get("bold", False)
        italic = blk.get("italic", False)

        if bold and italic:
            rl_font = fname + ("-BoldOblique" if fname == "Helvetica" else "-BoldItalic")
        elif bold:
            rl_font = fname + "-Bold"
        elif italic:
            rl_font = fname + ("-Oblique" if fname == "Helvetica" else "-Italic")
        else:
            rl_font = fname

        style = ParagraphStyle(
            name=f"s{blk['page']}_{blk['block_idx']}",
            fontName=rl_font,
            fontSize=float(blk.get("font_size", 11.0)),
            textColor=HexColor(blk.get("font_color", "#000000")),
            leading=float(blk.get("font_size", 11.0)) * 1.4,
            spaceAfter=4,
        )
        safe = (blk["edited_text"]
                .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                .replace("\n", "<br/>"))
        try:
            story.append(Paragraph(safe, style))
        except Exception:
            story.append(Paragraph(safe.encode("ascii","replace").decode(), style))
        story.append(Spacer(1, 3))

    if not story:
        from reportlab.lib.styles import getSampleStyleSheet
        story.append(Paragraph("(vacío)", getSampleStyleSheet()["Normal"]))

    doc.build(story)
    return buf.getvalue()


def blocks_to_docx(blocks: list, title: str = "Documento") -> bytes:
    from docx import Document as DocxDoc
    from docx.shared import Pt, Cm, RGBColor

    doc = DocxDoc()
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.5)
    doc.add_heading(title, level=0)

    cur_page = -1
    for blk in sorted(blocks, key=lambda b: (b["page"], b.get("y0", 0))):
        if blk["page"] != cur_page:
            if cur_page >= 0:
                doc.add_page_break()
            cur_page = blk["page"]
        for line in blk["edited_text"].split("\n"):
            if not line.strip():
                continue
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold   = blk.get("bold", False)
            run.italic = blk.get("italic", False)
            run.font.size = Pt(float(blk.get("font_size", 11.0)))
            c = blk.get("font_color", "#000000").lstrip("#")
            try:
                run.font.color.rgb = RGBColor(int(c[0:2],16), int(c[2:4],16), int(c[4:6],16))
            except Exception:
                pass

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_fake_blocks(text: str) -> list:
    return [
        {"page": 0, "block_idx": i, "original_text": p, "edited_text": p,
         "x0": 50.0, "y0": float(i*60), "x1": 550.0, "y1": float(i*60+50),
         "font_name": "Helvetica", "font_size": 11.0,
         "font_color": "#000000", "bold": False, "italic": False}
        for i, p in enumerate(text.split("\n\n")) if p.strip()
    ]


def validate_size(f, mb=50):
    return f.size <= mb * 1024 * 1024


# ══════════════════════════════════════════════════════════════════════════════
# TABS
# ══════════════════════════════════════════════════════════════════════════════
tab_editor, tab_converter, tab_ai = st.tabs([
    "✏️  Editar PDF",
    "🔄  Convertir",
    "✦  Asistente IA",
])


# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — EDITOR
# ══════════════════════════════════════════════════════════════════════════════
with tab_editor:
    st.markdown('<p class="sec-title">Editor estructurado de PDF</p>', unsafe_allow_html=True)
    st.markdown('<p class="sec-sub">Edita bloque a bloque · Preserva estructura original · Controla fuente, tamaño y color</p>', unsafe_allow_html=True)

    uploaded = st.file_uploader("📄 Sube tu PDF (máx. 50 MB)", type=["pdf"], key="editor_upload")

    if uploaded:
        if not validate_size(uploaded):
            st.markdown('<div class="err-box">❌ El archivo supera 50 MB.</div>', unsafe_allow_html=True)
        elif uploaded.name != st.session_state["pdf_filename"]:
            with st.spinner("Analizando estructura del PDF..."):
                raw = uploaded.read()
                blocks, npages, err = extract_blocks(raw)
            if err:
                st.markdown(f'<div class="err-box">❌ No se pudo leer el PDF: {err}</div>', unsafe_allow_html=True)
            else:
                st.session_state.update({
                    "blocks": blocks,
                    "pdf_bytes": raw,
                    "pdf_filename": uploaded.name,
                    "n_pages": npages,
                })
                st.markdown(
                    f'<div class="ok-box">✓ <strong>{uploaded.name}</strong> cargado · '
                    f'{npages} página(s) · {len(blocks)} bloques de texto</div>',
                    unsafe_allow_html=True
                )

    # ── Workspace ─────────────────────────────────────────────────────────────
    if st.session_state["blocks"]:
        st.markdown(
            f'<div class="info-bar">📄 <strong>{st.session_state["pdf_filename"]}</strong> · '
            f'{st.session_state["n_pages"]} págs. · {len(st.session_state["blocks"])} bloques</div>',
            unsafe_allow_html=True
        )

        col_main, col_side = st.columns([3, 1])

        # ── SIDE PANEL ────────────────────────────────────────────────────────
        with col_side:

            # Stats
            changed_count = sum(1 for b in st.session_state["blocks"]
                                if b["edited_text"] != b["original_text"])
            total_words   = sum(len(b["edited_text"].split()) for b in st.session_state["blocks"])
            st.metric("Bloques modificados", f"{changed_count} / {len(st.session_state['blocks'])}")
            st.metric("Palabras", f"{total_words:,}")

            st.markdown("---")

            # Search & Replace
            with st.expander("🔍 Buscar & Reemplazar"):
                find_t   = st.text_input("Buscar",         key="fin", placeholder="texto...")
                repl_t   = st.text_input("Reemplazar con", key="rep", placeholder="nuevo...")
                case_s   = st.checkbox("Distinguir mayúsculas", key="css")
                c1, c2   = st.columns(2)
                if c1.button("Reemplazar", key="brepl", use_container_width=True):
                    if find_t.strip():
                        n = 0
                        for b in st.session_state["blocks"]:
                            src = b["edited_text"]
                            if case_s:
                                n += src.count(find_t)
                                b["edited_text"] = src.replace(find_t, repl_t)
                            else:
                                n += len(re.findall(re.escape(find_t), src, re.IGNORECASE))
                                b["edited_text"] = re.sub(re.escape(find_t), repl_t, src, flags=re.IGNORECASE)
                        st.markdown(f'<div class="ok-box">✓ {n} reemplazos</div>', unsafe_allow_html=True)
                        st.rerun()
                if c2.button("Restaurar", key="brst", use_container_width=True):
                    for b in st.session_state["blocks"]:
                        b["edited_text"] = b["original_text"]
                    st.markdown('<div class="ok-box">✓ Texto restaurado</div>', unsafe_allow_html=True)
                    st.rerun()

            st.markdown("---")

            # Global font controls
            with st.expander("🎨 Fuente global (todos los bloques)", expanded=True):
                g_font  = st.selectbox("Tipo de letra", ["Helvetica","Times New Roman","Courier"], key="gf")
                g_size  = st.slider("Tamaño (pt)", 6, 48, 11, key="gs")
                g_color = st.color_picker("Color", "#000000", key="gc")
                gb_col, gi_col = st.columns(2)
                g_bold   = gb_col.checkbox("Negrita",  key="gbd")
                g_italic = gi_col.checkbox("Cursiva",  key="git")

                if st.button("✅ Aplicar a todo el PDF", key="bapply", use_container_width=True):
                    for b in st.session_state["blocks"]:
                        b["font_name"]  = g_font
                        b["font_size"]  = float(g_size)
                        b["font_color"] = g_color
                        b["bold"]       = g_bold
                        b["italic"]     = g_italic
                    st.markdown('<div class="ok-box">✓ Estilo aplicado a todos los bloques</div>', unsafe_allow_html=True)
                    st.rerun()

            st.markdown("---")

            # Downloads
            with st.expander("📥 Descargar resultado", expanded=True):
                base = st.session_state["pdf_filename"].replace(".pdf", "")

                try:
                    pdf_out = blocks_to_pdf(
                        st.session_state["blocks"],
                        st.session_state["pdf_bytes"]
                    )
                    st.download_button("📥 PDF editado", data=pdf_out,
                                       file_name=f"{base}_editado.pdf",
                                       mime="application/pdf",
                                       use_container_width=True, key="dl_pdf")
                except Exception as e:
                    st.markdown(f'<div class="err-box">PDF error: {e}</div>', unsafe_allow_html=True)

                txt_out = "\n\n".join(
                    f"[Página {b['page']+1} · Bloque {b['block_idx']+1}]\n{b['edited_text']}"
                    for b in st.session_state["blocks"]
                )
                st.download_button("📄 TXT", data=txt_out.encode("utf-8"),
                                   file_name=f"{base}_editado.txt", mime="text/plain",
                                   use_container_width=True, key="dl_txt")

                try:
                    st.download_button("📝 Word (DOCX)",
                                       data=blocks_to_docx(st.session_state["blocks"], base),
                                       file_name=f"{base}_editado.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       use_container_width=True, key="dl_docx")
                except Exception as e:
                    st.markdown(f'<div class="err-box">Word error: {e}</div>', unsafe_allow_html=True)

        # ── MAIN EDITOR AREA ──────────────────────────────────────────────────
        with col_main:
            pages = sorted(set(b["page"] for b in st.session_state["blocks"]))
            page_labels = ["Todas las páginas"] + [f"Página {p+1}" for p in pages]
            sel_page = st.selectbox("Mostrar:", page_labels, key="pfilt")
            vis_pages = pages if sel_page == "Todas las páginas" else [int(sel_page.split()[-1])-1]

            st.caption("✏️ Cada bloque es editable independientemente. Expande el encabezado del bloque para cambiar su fuente.")

            cur_p = -1
            for idx, blk in enumerate(st.session_state["blocks"]):
                if blk["page"] not in vis_pages:
                    continue

                if blk["page"] != cur_p:
                    cur_p = blk["page"]
                    st.markdown(
                        f'<div class="page-divider">─── Página {blk["page"]+1} ───</div>',
                        unsafe_allow_html=True
                    )

                edited_flag = blk["edited_text"] != blk["original_text"]
                tag = "✏️ " if edited_flag else ""

                # Font preview in expander title
                fpreview = (
                    f"{blk.get('font_name','Helvetica')} · "
                    f"{blk.get('font_size',11)}pt"
                    + (" · B" if blk.get("bold") else "")
                    + (" · I" if blk.get("italic") else "")
                    + (f" · {blk.get('font_color','')}" if blk.get("font_color","#000000") != "#000000" else "")
                )

                with st.expander(
                    f"{tag}Bloque {blk['block_idx']+1}  [{fpreview}]",
                    expanded=edited_flag,
                ):
                    # Per-block font controls
                    fa, fb_, fc_, fd_, fe_ = st.columns([2.5, 1.2, 1.2, 0.8, 0.8])

                    fonts_list = ["Helvetica", "Times New Roman", "Courier"]
                    cur_font_idx = fonts_list.index(blk.get("font_name","Helvetica")) if blk.get("font_name","Helvetica") in fonts_list else 0

                    new_font  = fa.selectbox("Fuente", fonts_list, index=cur_font_idx, key=f"f_{idx}")
                    new_size  = fb_.number_input("Tamaño", 6, 72,
                                                 int(blk.get("font_size", 11)), key=f"s_{idx}")
                    new_color = fc_.color_picker("Color", blk.get("font_color","#000000"), key=f"c_{idx}")
                    new_bold  = fd_.checkbox("B", blk.get("bold",  False), key=f"b_{idx}")
                    new_ital  = fe_.checkbox("I", blk.get("italic",False), key=f"i_{idx}")

                    # Update font settings
                    st.session_state["blocks"][idx].update({
                        "font_name":  new_font,
                        "font_size":  float(new_size),
                        "font_color": new_color,
                        "bold":       new_bold,
                        "italic":     new_ital,
                    })

                # Editable text (outside expander so always visible)
                n_lines = max(2, len(blk["edited_text"].split("\n")))
                new_text = st.text_area(
                    f"blk_text_{idx}",
                    value=blk["edited_text"],
                    height=min(300, max(70, n_lines * 26 + 30)),
                    key=f"t_{idx}",
                    label_visibility="collapsed",
                )
                if new_text != st.session_state["blocks"][idx]["edited_text"]:
                    st.session_state["blocks"][idx]["edited_text"] = new_text

    else:
        st.markdown("""
        <div style="border:2px dashed #888;border-radius:4px;padding:80px 36px;
                    text-align:center;background:#fff;margin-top:16px;">
          <div style="font-size:3.5rem;margin-bottom:14px;">📄</div>
          <div style="font-family:'Fraunces',serif;font-size:1.15rem;font-weight:300;color:#111;">
            Sube un PDF para editar bloque a bloque
          </div>
          <div style="font-size:0.65rem;letter-spacing:0.12em;text-transform:uppercase;
                      color:#666;margin-top:10px;line-height:2.2;">
            Preserva la estructura original del PDF<br>
            Controla fuente · tamaño · color · negrita · cursiva por bloque<br>
            Buscar y reemplazar en todo el documento<br>
            Exporta como PDF · TXT · Word
          </div>
        </div>
        """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — CONVERTER
# ══════════════════════════════════════════════════════════════════════════════
with tab_converter:
    st.markdown('<p class="sec-title">Conversor de formatos</p>', unsafe_allow_html=True)
    st.markdown('<p class="sec-sub">PDF · JPG · TXT · Word — conversiones en ambas direcciones</p>', unsafe_allow_html=True)

    st.markdown("""
    <div style="display:grid;grid-template-columns:repeat(4,1fr);gap:10px;margin-bottom:22px;">
      <div style="background:#fff;border:1.5px solid #111;border-radius:4px;padding:14px;text-align:center;">
        <div style="font-family:'Fraunces',serif;font-size:1.1rem;color:#c47b28;">PDF</div>
        <div style="color:#888;font-size:1rem;">↕</div>
        <div style="font-family:'Fraunces',serif;font-size:1.1rem;">TXT</div>
        <div style="font-size:0.58rem;text-transform:uppercase;letter-spacing:0.1em;color:#888;margin-top:4px;">texto plano</div>
      </div>
      <div style="background:#fff;border:1.5px solid #111;border-radius:4px;padding:14px;text-align:center;">
        <div style="font-family:'Fraunces',serif;font-size:1.1rem;color:#c47b28;">PDF</div>
        <div style="color:#888;font-size:1rem;">↓</div>
        <div style="font-family:'Fraunces',serif;font-size:1.1rem;">JPG</div>
        <div style="font-size:0.58rem;text-transform:uppercase;letter-spacing:0.1em;color:#888;margin-top:4px;">cada página</div>
      </div>
      <div style="background:#fff;border:1.5px solid #111;border-radius:4px;padding:14px;text-align:center;">
        <div style="font-family:'Fraunces',serif;font-size:1.1rem;color:#c47b28;">PDF</div>
        <div style="color:#888;font-size:1rem;">↕</div>
        <div style="font-family:'Fraunces',serif;font-size:1.1rem;">Word</div>
        <div style="font-size:0.58rem;text-transform:uppercase;letter-spacing:0.1em;color:#888;margin-top:4px;">editable</div>
      </div>
      <div style="background:#fff;border:1.5px solid #111;border-radius:4px;padding:14px;text-align:center;">
        <div style="font-family:'Fraunces',serif;font-size:1.1rem;color:#c47b28;">IMG</div>
        <div style="color:#888;font-size:1rem;">↑</div>
        <div style="font-family:'Fraunces',serif;font-size:1.1rem;">PDF</div>
        <div style="font-size:0.58rem;text-transform:uppercase;letter-spacing:0.1em;color:#888;margin-top:4px;">imagen→doc</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    cu, cf = st.columns([1, 1])
    with cu:
        conv_file = st.file_uploader("📁 Sube el archivo",
                                     type=["pdf","jpg","jpeg","png","txt","docx"],
                                     key="conv_upload")
        if conv_file:
            if not validate_size(conv_file, 100):
                st.markdown('<div class="err-box">❌ Supera 100 MB.</div>', unsafe_allow_html=True)
                conv_file = None
            else:
                ext = Path(conv_file.name).suffix.lower().lstrip(".")
                icons = {"pdf":"📄","jpg":"🖼️","jpeg":"🖼️","png":"🖼️","txt":"📝","docx":"📋"}
                st.markdown(
                    f'<div class="info-bar">{icons.get(ext,"📁")} <strong>{conv_file.name}</strong> '
                    f'· {conv_file.size/1024:.1f} KB · {ext.upper()}</div>',
                    unsafe_allow_html=True)

    with cf:
        if conv_file:
            ext  = Path(conv_file.name).suffix.lower().lstrip(".")
            opts = {
                "pdf":  ["TXT — Texto plano","JPG — Imagen por página","Word — DOCX editable"],
                "jpg":  ["PDF — Documento"],
                "jpeg": ["PDF — Documento"],
                "png":  ["PDF — Documento","JPG — Convertir formato"],
                "txt":  ["PDF — Documento","Word — DOCX editable"],
                "docx": ["PDF — Documento","TXT — Texto plano"],
            }
            target = st.selectbox("Convertir a →", opts.get(ext, ["TXT — Texto plano"]), key="ctarget")
            go = st.button("🔄 Convertir", key="bconv", use_container_width=True)

            if go:
                fb   = conv_file.read()
                base = Path(conv_file.name).stem
                res  = []
                with st.spinner("Convirtiendo..."):
                    try:
                        if ext == "pdf" and "TXT" in target:
                            blks, _, _ = extract_blocks(fb)
                            txt = "\n\n".join(f"[Pág {b['page']+1}]\n{b['original_text']}" for b in blks)
                            res.append((txt.encode("utf-8"), f"{base}.txt", "text/plain"))

                        elif ext == "pdf" and "JPG" in target:
                            import fitz
                            doc = fitz.open(stream=fb, filetype="pdf")
                            for i, pg in enumerate(doc):
                                pix = pg.get_pixmap(matrix=fitz.Matrix(2,2))
                                res.append((pix.tobytes("jpeg"), f"{base}_p{i+1}.jpg", "image/jpeg"))

                        elif ext == "pdf" and "Word" in target:
                            blks, _, _ = extract_blocks(fb)
                            res.append((blocks_to_docx(blks, base), f"{base}.docx",
                                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))

                        elif ext in ("jpg","jpeg","png") and "PDF" in target:
                            from reportlab.pdfgen import canvas as rlc
                            from reportlab.lib.utils import ImageReader
                            from PIL import Image as PILImage
                            img = PILImage.open(io.BytesIO(fb))
                            if img.mode in ("RGBA","P"):
                                img = img.convert("RGB")
                            w, h = img.size
                            ib = io.BytesIO(); img.save(ib,"JPEG",quality=95); ib.seek(0)
                            pb = io.BytesIO()
                            c = rlc.Canvas(pb, pagesize=(w,h))
                            c.drawImage(ImageReader(ib), 0, 0, w, h); c.save()
                            res.append((pb.getvalue(), f"{base}.pdf", "application/pdf"))

                        elif ext == "png" and "JPG" in target:
                            from PIL import Image as PILImage
                            img = PILImage.open(io.BytesIO(fb)).convert("RGB")
                            pb = io.BytesIO(); img.save(pb,"JPEG",quality=92)
                            res.append((pb.getvalue(), f"{base}.jpg", "image/jpeg"))

                        elif ext == "txt" and "PDF" in target:
                            text = fb.decode("utf-8", errors="replace")
                            res.append((_blocks_to_pdf_rl(make_fake_blocks(text)),
                                        f"{base}.pdf", "application/pdf"))

                        elif ext == "txt" and "Word" in target:
                            text = fb.decode("utf-8", errors="replace")
                            res.append((blocks_to_docx(make_fake_blocks(text), base),
                                        f"{base}.docx",
                                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))

                        elif ext == "docx" and "PDF" in target:
                            from docx import Document as DocxDoc
                            dw = DocxDoc(io.BytesIO(fb))
                            t  = "\n\n".join(p.text for p in dw.paragraphs if p.text.strip())
                            res.append((_blocks_to_pdf_rl(make_fake_blocks(t)),
                                        f"{base}.pdf", "application/pdf"))

                        elif ext == "docx" and "TXT" in target:
                            from docx import Document as DocxDoc
                            dw = DocxDoc(io.BytesIO(fb))
                            t  = "\n".join(p.text for p in dw.paragraphs)
                            res.append((t.encode("utf-8"), f"{base}.txt", "text/plain"))

                        st.session_state["conv_results"] = res

                    except Exception as e:
                        st.markdown(f'<div class="err-box">❌ Error: {e}</div>', unsafe_allow_html=True)
        else:
            st.info("Sube un archivo para ver las opciones.")

    if st.session_state["conv_results"]:
        st.markdown("---")
        res = st.session_state["conv_results"]
        st.markdown(f'<div class="ok-box">✓ {len(res)} archivo(s) listo(s)</div>', unsafe_allow_html=True)
        if len(res) == 1:
            d, fn, mime = res[0]
            st.download_button(f"📥 Descargar {fn}", data=d, file_name=fn, mime=mime,
                               use_container_width=True, key="dlc1")
        else:
            zb = io.BytesIO()
            with zipfile.ZipFile(zb, "w", zipfile.ZIP_DEFLATED) as zf:
                for d, fn, _ in res:
                    zf.writestr(fn, d)
            st.download_button("📦 Descargar todo (ZIP)", data=zb.getvalue(),
                               file_name="convertidos.zip", mime="application/zip",
                               use_container_width=True, key="dlzip")
            for i, (d, fn, mime) in enumerate(res):
                st.download_button(f"  └ {fn}", data=d, file_name=fn, mime=mime,
                                   key=f"dlci{i}", use_container_width=True)


# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — AI ASSISTANT
# ══════════════════════════════════════════════════════════════════════════════
with tab_ai:
    st.markdown('<p class="sec-title">Asistente IA — Claude</p>', unsafe_allow_html=True)
    st.markdown('<p class="sec-sub">Procesa, corrige, resume y transforma el documento con inteligencia artificial</p>', unsafe_allow_html=True)

    with st.expander("⚙️ Clave API de Anthropic", expanded=not bool(st.session_state["api_key"])):
        api_in = st.text_input("Clave API (sk-ant-...)", type="password",
                               value=st.session_state["api_key"], key="api_in")
        if api_in != st.session_state["api_key"]:
            st.session_state["api_key"] = api_in
        if st.session_state["api_key"]:
            st.markdown('<div class="ok-box">✓ Clave configurada</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="warn-box">⚠️ Obtén tu clave en <strong>console.anthropic.com</strong></div>', unsafe_allow_html=True)

    al, ar = st.columns([1, 1])

    with al:
        mode = st.radio("Fuente:", ["Pegar manualmente","Usar PDF del editor"],
                        horizontal=True, key="aimode")

        if mode == "Usar PDF del editor":
            if st.session_state["blocks"]:
                ai_text = "\n\n".join(b["edited_text"] for b in st.session_state["blocks"])
                st.markdown(f'<div class="info-bar">📄 {st.session_state["pdf_filename"]} · {len(ai_text):,} chars</div>', unsafe_allow_html=True)
                st.text_area("Previa:", value=ai_text[:500]+"...", height=120, disabled=True, key="aiprev")
            else:
                st.markdown('<div class="warn-box">⚠️ Carga un PDF en el editor primero.</div>', unsafe_allow_html=True)
                ai_text = ""
        else:
            ai_text = st.text_area("Pega el texto:", height=180, key="aimanual",
                                   placeholder="Pega aquí el contenido...")

        st.markdown("**Instrucciones rápidas:**")
        quick = {
            "✓ Corregir ortografía": "Corrige todos los errores ortográficos y gramaticales sin cambiar el contenido.",
            "📝 Resumir (5 pts)":    "Resume en 5 puntos clave, en español.",
            "💼 Tono formal":        "Reescribe con tono profesional sin cambiar información.",
            "😊 Tono amigable":      "Reescribe con tono cercano y amigable.",
            "🌐 → Inglés":          "Traduce al inglés manteniendo el formato.",
            "🌐 → Español":         "Traduce al español manteniendo el formato.",
            "📋 Puntos clave":       "Extrae los puntos más importantes como lista numerada.",
            "✏️ Ampliar":           "Amplía el texto con detalles relevantes.",
        }
        qs = st.columns(2)
        for i, (lbl, ptxt) in enumerate(quick.items()):
            if qs[i%2].button(lbl, key=f"qai{i}", use_container_width=True):
                st.session_state["ai_prompt_v"] = ptxt
                st.rerun()

        prompt = st.text_area("O escribe tu instrucción:",
                              value=st.session_state.get("ai_prompt_v",""),
                              height=90, key="aipromptw",
                              placeholder='Ej: "Cambia todos los precios de USD a EUR"')
        if prompt != st.session_state.get("ai_prompt_v",""):
            st.session_state["ai_prompt_v"] = prompt

        run_ai = st.button("✦ Procesar con Claude", key="airun", use_container_width=True)

    with ar:
        st.markdown("**Resultado:**")

        if run_ai:
            final = ai_text if mode == "Pegar manualmente" else "\n\n".join(b["edited_text"] for b in st.session_state["blocks"])
            pv    = st.session_state.get("ai_prompt_v","")

            if not st.session_state["api_key"].strip():
                st.markdown('<div class="err-box">❌ Configura la clave API.</div>', unsafe_allow_html=True)
            elif not final.strip():
                st.markdown('<div class="err-box">❌ No hay texto para procesar.</div>', unsafe_allow_html=True)
            elif not pv.strip():
                st.markdown('<div class="err-box">❌ Elige o escribe una instrucción.</div>', unsafe_allow_html=True)
            else:
                with st.spinner("✦ Claude procesando..."):
                    try:
                        import anthropic
                        client = anthropic.Anthropic(api_key=st.session_state["api_key"].strip())
                        msg = client.messages.create(
                            model="claude-sonnet-4-20250514", max_tokens=4096,
                            messages=[{"role":"user",
                                       "content":f"Instrucción: {pv}\n\n---\n\n{final[:100000]}"}]
                        )
                        st.session_state["ai_result"] = msg.content[0].text
                        st.markdown('<div class="ok-box">✓ Procesado.</div>', unsafe_allow_html=True)
                    except Exception as e:
                        em = str(e)
                        label = "Clave inválida" if "401" in em else "Límite de velocidad" if "429" in em else em
                        st.markdown(f'<div class="err-box">❌ {label}</div>', unsafe_allow_html=True)

        result = st.session_state.get("ai_result","")
        if result:
            st.text_area("Resultado (editable):", value=result, height=320,
                         key="airesedit", label_visibility="visible")
            st.markdown("**Exportar:**")
            e1, e2, e3, e4 = st.columns(4)
            e1.download_button("📄 TXT", data=result.encode("utf-8"),
                               file_name="resultado_ia.txt", mime="text/plain",
                               key="dlaitxt", use_container_width=True)
            try:
                fb = make_fake_blocks(result)
                e2.download_button("📥 PDF", data=_blocks_to_pdf_rl(fb),
                                   file_name="resultado_ia.pdf", mime="application/pdf",
                                   key="dlaipdf", use_container_width=True)
                e3.download_button("📝 Word", data=blocks_to_docx(fb,"Resultado IA"),
                                   file_name="resultado_ia.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   key="dlaidocx", use_container_width=True)
            except Exception:
                pass

            if e4.button("✏️ Al editor", key="aitoeditor", use_container_width=True):
                nb = make_fake_blocks(result)
                st.session_state.update({
                    "blocks": nb,
                    "pdf_filename": "resultado_ia.pdf",
                    "n_pages": 1,
                    "pdf_bytes": None,
                })
                st.markdown('<div class="ok-box">✓ Enviado al editor.</div>', unsafe_allow_html=True)

            if st.button("🗑 Limpiar", key="aiclear"):
                st.session_state["ai_result"] = ""
                st.rerun()
        else:
            st.markdown("""
            <div style="border:1px dashed #bbb;border-radius:3px;padding:80px 24px;
                        text-align:center;color:#888;font-size:0.73rem;letter-spacing:0.08em;">
              ✦ El resultado de Claude aparecerá aquí
            </div>""", unsafe_allow_html=True)

# ─── FOOTER ───────────────────────────────────────────────────────────────────
st.markdown("---")
st.markdown("""
<div style="text-align:center;font-size:0.6rem;letter-spacing:0.12em;
            text-transform:uppercase;color:#888;padding:6px 0;">
  PDF Studio · Streamlit · IA por Claude — Anthropic
</div>""", unsafe_allow_html=True)
